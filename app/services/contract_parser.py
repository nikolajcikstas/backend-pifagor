"""
Парсинг .docx-договоров, загружаемых админом во вкладке «Договоры».

Извлекает:
- ФИО клиента (Заказчика) из шапки договора — для привязки к ученику
- даты начала/окончания оказания услуг
- общую стоимость услуг
- график платежей (номер, сумма, срок оплаты)
- контактные данные родителя из блока реквизитов в конце договора
  (ФИО, телефон, email, город/улица/дом — индекс и квартира используются
  только для очистки адреса и никуда не сохраняются)
"""
import re
from datetime import date
from typing import Optional

import docx

MONTHS = {
    "января": 1, "февраля": 2, "марта": 3, "апреля": 4, "мая": 5, "июня": 6,
    "июля": 7, "августа": 8, "сентября": 9, "октября": 10, "ноября": 11, "декабря": 12,
}

ORDINALS = {
    "первый": 1, "второй": 2, "третий": 3, "четвёртый": 4, "четвертый": 4, "пятый": 5,
    "шестой": 6, "седьмой": 7, "восьмой": 8, "девятый": 9, "десятый": 10,
    "одиннадцатый": 11, "двенадцатый": 12, "тринадцатый": 13, "четырнадцатый": 14,
    "пятнадцатый": 15, "шестнадцатый": 16, "семнадцатый": 17, "восемнадцатый": 18,
    "девятнадцатый": 19, "двадцатый": 20,
    "двадцать первый": 21, "двадцать второй": 22, "двадцать третий": 23,
    "двадцать четвёртый": 24, "двадцать четвертый": 24, "двадцать пятый": 25,
    "двадцать шестой": 26, "двадцать седьмой": 27, "двадцать восьмой": 28,
    "двадцать девятый": 29, "двадцать тридцатый": 30, "тридцатый": 30,
    "тридцать первый": 31, "тридцать второй": 32, "тридцать третий": 33,
    "тридцать четвёртый": 34, "тридцать четвертый": 34, "тридцать пятый": 35,
    "тридцать шестой": 36, "тридцать седьмой": 37, "тридцать восьмой": 38,
    "тридцать девятый": 39, "сороковой": 40,
}


def _parse_ru_date(text: str) -> Optional[date]:
    text = text.strip().rstrip("г.").strip()
    m = re.match(r"(\d{1,2})\s+(\S+)\s+(\d{4})", text)
    if not m:
        return None
    day, month_name, year = m.groups()
    month = MONTHS.get(month_name.lower())
    if not month:
        return None
    try:
        return date(int(year), month, int(day))
    except ValueError:
        return None


def _parse_amount(text: str) -> Optional[float]:
    cleaned = re.sub(r"[^0-9,.]", "", text)
    cleaned = cleaned.replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return None


def parse_contract_docx(file_path: str) -> dict:
    d = docx.Document(file_path)
    full_text = "\n".join(p.text for p in d.paragraphs)

    result: dict = {
        "contract_number": None,
        "client_name_header": None,
        "start_date": None,
        "end_date": None,
        "total_amount": None,
        "payments": [],
        "parent_full_name": None,
        "parent_phone": None,
        "parent_email": None,
        "city": None,
        "street": None,
        "house": None,
        "parse_warnings": [],
    }

    m = re.search(r"ДОГОВОР\s*№\s*(\S+)", full_text)
    if m:
        result["contract_number"] = m.group(1).strip()

    m = re.search(r"с одной стороны,\s*(.+?),\s*именуем", full_text)
    if m:
        result["client_name_header"] = re.sub(r"\s+", " ", m.group(1)).strip()
    else:
        result["parse_warnings"].append("Не удалось найти ФИО заказчика в шапке договора")

    m = re.search(r"Начало оказания Услуг\s*[–-]\s*(.+?)(?:г\.|\n)", full_text)
    result["start_date"] = _parse_ru_date(m.group(1)) if m else None
    m = re.search(r"Дата окончания Услуг\s*[–-]\s*(.+?)(?:г\.|\n)", full_text)
    result["end_date"] = _parse_ru_date(m.group(1)) if m else None
    if result["start_date"] and result["end_date"] and result["end_date"] < result["start_date"]:
        result["parse_warnings"].append(
            "Дата окончания раньше даты начала — вероятно, опечатка в тексте договора"
        )

    m = re.search(r"Стоимость Услуг составляет\s*([\d\s\xa0\u202f]+,\d+)", full_text)
    result["total_amount"] = _parse_amount(m.group(1)) if m else None

    payments = []
    pattern = re.compile(
        r"Заказчик вносит (\S+(?:\s\S+)?) плат[её]ж\s*([\d\s\xa0\u202f]+,\d+).*?"
        r"не позднее\s*(\d{1,2}\s+\S+\s+\d{4})\s*года",
        re.IGNORECASE,
    )
    for match in pattern.finditer(full_text):
        ordinal_word, amount_text, date_text = match.groups()
        ordinal_word_clean = ordinal_word.strip().lower()
        number = ORDINALS.get(ordinal_word_clean)
        due_date = _parse_ru_date(date_text)
        amount = _parse_amount(amount_text)
        if number is None or due_date is None or amount is None:
            result["parse_warnings"].append(f"Не разобрана строка платежа: «{ordinal_word}»")
            continue
        payments.append({"number": number, "amount": amount, "due_date": due_date.isoformat()})
    result["payments"] = payments

    # Некоторые договоры вообще не содержат пронумерованный график платежей —
    # оплата идёт по факту, после каждого занятия, либо ежемесячно к
    # определённому числу. Это не ошибка парсинга — просто другой тип
    # договора, для него формула "отработать N занятий" не применяется.
    if payments:
        result["payment_mode"] = "scheduled"
    elif re.search(r"в день проведения (?:онлайн-встреч|занят)", full_text, re.IGNORECASE):
        result["payment_mode"] = "per_lesson"
    elif re.search(r"\d{1,2}-го числа каждого месяца", full_text, re.IGNORECASE):
        result["payment_mode"] = "monthly"
    else:
        result["payment_mode"] = "unknown"
        result["parse_warnings"].append("Не найдено ни одного платежа в графике")

    if not d.tables:
        result["parse_warnings"].append("В документе не найдена таблица с реквизитами сторон")
        return result

    table = d.tables[-1]
    if len(table.rows) < 3:
        result["parse_warnings"].append("Таблица реквизитов имеет неожиданную структуру")
        return result

    cell_text = table.rows[2].cells[0].text
    lines = [l.strip() for l in cell_text.split("\n") if l.strip()]

    phone_m = re.search(r"(\+375\d{9}|\b80\d{9}\b)", cell_text)
    email_m = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", cell_text)
    result["parent_phone"] = phone_m.group(1) if phone_m else None
    result["parent_email"] = email_m.group(0) if email_m else None
    if not result["parent_phone"]:
        result["parse_warnings"].append("Не найден номер телефона в блоке реквизитов")
    if not result["parent_email"]:
        result["parse_warnings"].append("Не найден email в блоке реквизитов")

    if lines:
        # ФИО обычно первая строка, но иногда телефон приклеен туда же без
        # переноса строки — вырезаем его, если он там нашёлся.
        name_line = lines[0]
        if result["parent_phone"]:
            name_line = name_line.replace(result["parent_phone"], "")
        result["parent_full_name"] = re.sub(r"\s+", " ", name_line).strip(" .")
    else:
        result["parse_warnings"].append("Не удалось найти ФИО в блоке реквизитов")

    address_line = None
    for l in lines[1:]:
        if result["parent_phone"] and result["parent_phone"] in l:
            continue
        if result["parent_email"] and result["parent_email"] in l:
            continue
        if l.startswith("_") or "Б.П." in l:
            continue
        address_line = l
        break

    if address_line:
        addr = address_line
        index_m = re.search(r"\b(\d{6})\b", addr)
        if index_m:
            addr = addr.replace(index_m.group(1), "")

        flat_m = re.search(r"кв\.?\s*(\d+)", addr, re.IGNORECASE)
        if flat_m:
            addr = addr[:flat_m.start()] + addr[flat_m.end():]

        addr = re.sub(r"^[,.\s]+|[,.\s]+$", "", addr)
        parts = [p.strip() for p in re.split(r"[,]", addr) if p.strip()]

        # Тип населённого пункта (город/деревня/агрогородок/посёлок и т.п.) —
        # приравниваем к городу. Точка после сокращения ОБЯЗАТЕЛЬНА (кроме
        # полных слов) — иначе "п" без точки ложно совпадает с началом любого
        # слова на "П" (например «Платонова»). После сокращения обязательно
        # должна идти буква, а не цифра — иначе это «д.33» (номер дома), а не
        # «д. Тарасово» (деревня).
        SETTLEMENT_PREFIX = re.compile(
            r"^(?:г\.|гор\.|город|д\.|дер\.|деревня|аг\.|агрогородок|п\.|пос\.|посёлок|поселок|гп\.)"
            r"\s*([А-Яа-яЁё].*)$",
        )
        # Область/район — не город и не улица, отбрасываем как шум.
        NOISE_ONLY = re.compile(
            r"^.+\s+р-?н\.?$|^.+\s+район$|^.+\s+обл\.?$|^.+\s+область$", re.IGNORECASE
        )

        city = None
        street_parts = []
        for part in parts:
            if NOISE_ONLY.match(part):
                continue
            sm = SETTLEMENT_PREFIX.match(part)
            if sm and city is None:
                city = sm.group(1).strip(" .")
                continue
            street_parts.append(part)

        # Если явного "г./д./аг." нигде не было, но самый первый оставшийся
        # сегмент — это просто голое название города без номеров (например
        # "Минск" без "г."), и после него есть ещё сегменты с улицей/домом —
        # считаем его городом.
        if city is None and len(street_parts) > 1 and not any(ch.isdigit() for ch in street_parts[0]):
            city = street_parts.pop(0).strip(" .")

        # Адрес без единой запятой ("Минская обл. г.Дзержинск ул.Кооперативная 44") —
        # ищем населённый пункт где угодно внутри строки, а не только в начале.
        if city is None and len(street_parts) == 1:
            inner = re.search(
                r"(?:^|\s)(?:г\.|гор\.|город|д\.|дер\.|деревня|аг\.|агрогородок|п\.|пос\.|посёлок|поселок|гп\.)"
                r"\s*([А-Яа-яЁё]+)",
                street_parts[0],
            )
            if inner:
                city = inner.group(1).strip(" .")
                street_parts[0] = street_parts[0][inner.end():].strip(" .,")

        # Голый город без "г." и без запятой перед улицей ("Минск ул. Пономарева 9-3").
        if city is None and len(street_parts) == 1:
            bare_m = re.match(r"^([А-ЯЁ][а-яё]+)\s+(?:ул\.|улица|пр-т|пр\.|проспект)", street_parts[0])
            if bare_m:
                city = bare_m.group(1)
                street_parts[0] = street_parts[0][bare_m.end(1):].strip(" .,")

        street = None
        house = None
        if street_parts:
            last = street_parts[-1]
            # "д.33" — это номер дома ("дом"), а не населённый пункт; убираем
            # такой префикс перед тем, как отделять номер дома от улицы.
            last = re.sub(r"^д\.?\s*(?=\d)", "", last, flags=re.IGNORECASE)
            hm = re.search(r"(\d+(?:[-/]\d+)*[а-яА-Я]?)\s*$", last)
            if hm:
                house = hm.group(1)
                street = last[:hm.start()].strip(" .")
            else:
                street = last
            if len(street_parts) > 1:
                street = ", ".join(street_parts[:-1]) + (f", {street}" if street else "")

        result["city"] = city
        result["street"] = street
        result["house"] = house
        if not city:
            result["parse_warnings"].append(
                "Город не распознан однозначно — проверьте и поправьте вручную"
            )
    else:
        result["parse_warnings"].append("Не удалось найти строку адреса в блоке реквизитов")

    return result
